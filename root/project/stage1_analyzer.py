import os
import re
import sys
import json

from pathlib import Path
from datetime import datetime
from zoneinfo import ZoneInfo

import pytz
from babel import Locale
from openai import OpenAI


class TimeZoneResolver:
    """
    국가명 -> ISO 국가 코드 -> IANA TimeZone -> 현재 UTC Offset 계산
    """

    def __init__(self):
        korean_locale = Locale("ko")

        self.name_to_iso = {
            name: code
            for code, name in korean_locale.territories.items()
            if len(code) == 2
        }

        # 자주 사용하는 이형 표현
        self.name_to_iso.update({
            "한국": "KR",
            "대한민국": "KR",
            "미국": "US",
            "영국": "GB",
        })

    # ---------------------------------------------------------
    # 국가명 -> ISO 국가 코드
    # ---------------------------------------------------------
    def resolve_country_code(self, country_name):

        if not country_name:
            return None

        country_name = country_name.strip()

        # 정확히 일치
        if country_name in self.name_to_iso:
            return self.name_to_iso[country_name]

        # 부분 일치
        for name, code in self.name_to_iso.items():
            if country_name in name or name in country_name:
                return code

        return None

    # ---------------------------------------------------------
    # ISO 코드 -> 대표 TimeZone
    # ---------------------------------------------------------
    def get_primary_timezone(self, iso_code):

        tz_list = pytz.country_timezones.get(iso_code)

        if not tz_list:
            return None, 0

        # 다중 TimeZone 국가의 경우 첫 번째 대표값 사용
        return tz_list[0], len(tz_list)

    # ---------------------------------------------------------
    # 현재 시점 UTC Offset
    # DST 자동 반영
    # ---------------------------------------------------------
    @staticmethod
    def get_utc_offset_hours(tz_name):

        tz = ZoneInfo(tz_name)
        now = datetime.now(tz)

        offset = now.utcoffset()

        if offset is None:
            return 0

        return offset.total_seconds() / 3600


class OpenAIFileAnalyzer:

    BUSINESS_START = 9
    BUSINESS_END = 18

    # ---------------------------------------------------------
    # 초기화
    # ---------------------------------------------------------
    def __init__(self, input_file_path=None):

        self.content = None
        self.result = None
        self.time_zone_info = None

        self.tz_resolver = TimeZoneResolver()

        # =====================================================
        # 기준 경로
        # 현재 Python 파일이 있는 디렉터리
        # =====================================================
        self.base_path = Path(__file__).resolve().parent

        # =====================================================
        # 입력 파일
        #
        # 실행 시 파일 경로를 넘기면 해당 파일 사용
        #
        # 예:
        # python3 stage1_analyzer.py /root/files/meeting.pdf
        #
        # 경로를 안 넘기면:
        # 현재 폴더의 meeting_file.txt 사용
        # =====================================================
        if input_file_path:

            input_path = Path(input_file_path)

            if not input_path.is_absolute():
                input_path = Path.cwd() / input_path

            self.analyze_file_path = input_path.resolve()

        else:

            self.analyze_file_path = (
                self.base_path / "meeting_file.txt"
            )

        # =====================================================
        # 결과 파일
        # =====================================================
        self.result_path = (
            self.base_path / "openai_file_result.json"
        )

        # =====================================================
        # OpenAI 설정
        # =====================================================
        self.model_name = "gpt-4o-mini"

        self.api_key = self.load_api_key()

        self.client = OpenAI(
            api_key=self.api_key
        )

        # =====================================================
        # SYSTEM PROMPT
        # =====================================================
        self.SYSTEM_PROMPT = """
당신은 글로벌 비즈니스 협업을 지원하는 AI 분석가입니다.

입력으로
[사용자 정보],
[파트너 정보],
[협업 상황],
[회의 진행 내용],
[시차 정보 - 자동 계산됨]
이 주어집니다.

이 단계의 목적은 회의 내용을 정확하게 정리·요약하고,
협업에 필요한 문화적 인사이트를 제공하는 것입니다.

담당자별 To-Do 항목 추출은 이 단계에서 다루지 않습니다.


중요한 원칙:

1. cultural_analysis와 risk_notes는 일반적인 국가 상식이 아니라,
반드시 [회의 진행 내용]에서 실제로 관찰된 발언이나 행동을
근거로 작성하세요.

예:
"직접적 커뮤니케이션을 선호함" (X)

"Muller가 견적서의 근거 자료를 구체적으로 요청한 점에서
직접적이고 근거 중심의 소통 성향이 드러남" (O)


2. communication_style의 directness와 formality는
반드시 짧은 단어/구 형태로만 작성하세요.

예:
"직접적"
"격식 있는"

문장으로 서술하지 마세요.

판단 근거는 evidence 필드에서만 서술하세요.

단, tone_recommendation은 이 제약과 무관하게
실행 가능한 조언을 한두 문장으로 작성하세요.


3. 근거가 명확하지 않으면 절대로 해석을
억지로 만들어내지 마세요.

확실한 근거를 찾을 수 없는 경우,
해당 항목의 개수를 줄이거나 생략해도 됩니다.

추측성 문장보다 적은 수의 확실한 문장이 낫습니다.


4. time_zone_analysis는 절대 직접 계산하지 마세요.

[시차 정보 - 자동 계산됨] 블록에 이미 계산되어
제공된 값을 그대로 인용해서 사용하세요.


5. meeting_summary는 다음 두 단계로 작성하세요.

1단계(추출):
회의에서 언급된 모든 논의 주제를 먼저 전부 확인하세요.
사소해 보여도 화제로 오른 것은 전부 포함하세요.

2단계(분류):
각 주제를 다음 두 가지로 분류하세요.

- agreed
- disputed_or_pending

판단 기준:

구체적인 날짜·수치·행동이 확정되고
상대방이 동의를 표했으면 agreed입니다.

문제 제기만 되고 결론이 나지 않았거나
추가 확인이 필요하면 disputed_or_pending입니다.

같은 주제가 회의 중 여러 번 언급되면서
상황이 변경된 경우에는
회의 내 가장 마지막 발언을 기준으로
최종 상태만 판단하세요.

동일한 주제를 agreed와 disputed_or_pending
양쪽에 중복해서 넣지 마세요.


6. 최종 출력 전에
[회의 진행 내용]을 처음부터 끝까지 다시 확인하세요.

각 발화에서 언급된 안건이 meeting_summary에
모두 반영되었는지 확인하고,

동일한 주제가 중복 분류되지 않았는지 검토하세요.


출력은 반드시 아래 JSON 형식으로만 작성하세요.

다른 설명이나 텍스트는 절대 포함하지 마세요.


{
  "cultural_analysis": [
    "...",
    "..."
  ],

  "communication_style": {
    "directness": "...",
    "formality": "...",
    "evidence": "...",
    "tone_recommendation": "..."
  },

  "time_zone_analysis": {
    "time_difference": "...",
    "overlap_hours": "...",
    "recommended_meeting_time": "..."
  },

  "meeting_summary": {
    "agreed": [
      "...",
      "..."
    ],

    "disputed_or_pending": [
      "...",
      "..."
    ]
  },

  "risk_notes": [
    "...",
    "..."
  ]
}
""".strip()

    # ---------------------------------------------------------
    # API KEY 불러오기
    # ---------------------------------------------------------
    def load_api_key(self):

        api_key = os.environ.get("OPENAI_API_KEY")

        if not api_key:

            raise ValueError(
                "[ERROR] OPENAI_API_KEY 환경변수가 설정되어 있지 않습니다."
            )

        print(
            "[INFO] OPENAI_API_KEY 환경변수에서 API Key를 불러왔습니다."
        )

        return api_key.strip()

    # ---------------------------------------------------------
    # 기존 결과 제거
    # ---------------------------------------------------------
    def clear_previous_result(self):

        self.result_path.unlink(
            missing_ok=True
        )

        print(
            "[INFO] 기존 1단계 결과 삭제 완료"
        )

    # ---------------------------------------------------------
    # TXT 파일 읽기
    # ---------------------------------------------------------
    def read_txt(self):

        with open(
            self.analyze_file_path,
            "r",
            encoding="utf-8"
        ) as f:

            return f.read()

    # ---------------------------------------------------------
    # PDF 파일 읽기
    # ---------------------------------------------------------
    def read_pdf(self):

        try:

            from pypdf import PdfReader

        except ImportError:

            raise ImportError(
                "PDF 처리를 위해 pypdf 설치가 필요합니다.\n"
                "pip install pypdf"
            )

        reader = PdfReader(
            str(self.analyze_file_path)
        )

        pages = []

        for page_number, page in enumerate(
            reader.pages,
            start=1
        ):

            text = page.extract_text()

            if text:
                pages.append(text)

        return "\n\n".join(pages)

    # ---------------------------------------------------------
    # DOCX 파일 읽기
    # ---------------------------------------------------------
    def read_docx(self):

        try:

            from docx import Document

        except ImportError:

            raise ImportError(
                "DOCX 처리를 위해 python-docx 설치가 필요합니다.\n"
                "pip install python-docx"
            )

        document = Document(
            str(self.analyze_file_path)
        )

        text_parts = []

        # 일반 문단
        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:
                text_parts.append(text)

        # 표 안의 텍스트도 추출
        for table in document.tables:

            for row in table.rows:

                cells = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]

                if cells:
                    text_parts.append(
                        " | ".join(cells)
                    )

        return "\n".join(text_parts)

    # ---------------------------------------------------------
    # 입력 파일 읽기
    # TXT / PDF / DOCX 지원
    # ---------------------------------------------------------
    def openfile(self):

        try:

            if not self.analyze_file_path.is_file():

                print(
                    f"[ERROR] 입력 파일이 존재하지 않습니다.\n"
                    f"{self.analyze_file_path}"
                )

                return False

            extension = (
                self.analyze_file_path
                .suffix
                .lower()
            )

            if extension == ".txt":

                self.content = self.read_txt()

            elif extension == ".pdf":

                self.content = self.read_pdf()

            elif extension == ".docx":

                self.content = self.read_docx()

            else:

                print(
                    "[ERROR] 지원하지 않는 파일 형식입니다.\n"
                    f"현재 파일: {extension}\n"
                    "지원 형식: .txt, .pdf, .docx"
                )

                return False

            if not self.content:
                print(
                    "[ERROR] 파일에서 내용을 읽지 못했습니다."
                )
                return False

            self.content = self.content.strip()

            if not self.content:

                print(
                    "[ERROR] 분석 대상 파일 내용이 비어 있습니다."
                )

                return False

            print(
                f"[INFO] 파일 불러오기 성공: "
                f"{self.analyze_file_path}"
            )

            print(
                f"[INFO] 파일 형식: {extension}"
            )

            print(
                f"[INFO] 추출 문자 수: "
                f"{len(self.content)}"
            )

            return True

        except Exception as e:

            print(
                f"[ERROR] 파일 읽기 실패: {e}"
            )

            return False

    # ---------------------------------------------------------
    # 사용자/파트너 정보에서 국가 추출
    # ---------------------------------------------------------
    def extract_country(self, section_name):

        if not self.content:
            return None

        # 해당 섹션만 추출
        section_pattern = (
            rf"\[{re.escape(section_name)}\]"
            rf"(.*?)(?=\n\s*\[|\Z)"
        )

        section_match = re.search(
            section_pattern,
            self.content,
            re.DOTALL
        )

        if not section_match:
            return None

        section_content = (
            section_match
            .group(1)
        )

        country_match = re.search(
            r"국가\s*:\s*([^\n\r]+)",
            section_content
        )

        if not country_match:
            return None

        return (
            country_match
            .group(1)
            .strip()
        )

    # ---------------------------------------------------------
    # 시간 표시
    # ---------------------------------------------------------
    @staticmethod
    def format_hour(hour):

        # 분 단위로 변환해서 floating point 문제 방지
        total_minutes = (
            round((hour % 24) * 60)
            % (24 * 60)
        )

        whole_hour = (
            total_minutes // 60
        )

        minutes = (
            total_minutes % 60
        )

        period = (
            "오전"
            if whole_hour < 12
            else "오후"
        )

        display_hour = (
            whole_hour % 12
        )

        if display_hour == 0:
            display_hour = 12

        if minutes == 0:

            return (
                f"{period} {display_hour}시"
            )

        return (
            f"{period} "
            f"{display_hour}시 "
            f"{minutes}분"
        )

    # ---------------------------------------------------------
    # 시차 숫자 표시
    # 5.0 -> 5
    # 5.5 -> 5.5
    # ---------------------------------------------------------
    @staticmethod
    def format_difference(value):

        if float(value).is_integer():

            return str(
                int(abs(value))
            )

        return (
            f"{abs(value):.2f}"
            .rstrip("0")
            .rstrip(".")
        )

    # ---------------------------------------------------------
    # 시차 및 근무시간 계산
    # ---------------------------------------------------------
    def calculate_time_zone_info(self):

        user_country = (
            self.extract_country(
                "사용자 정보"
            )
        )

        partner_country = (
            self.extract_country(
                "파트너 정보"
            )
        )

        user_iso = (
            self.tz_resolver
            .resolve_country_code(
                user_country
            )
        )

        partner_iso = (
            self.tz_resolver
            .resolve_country_code(
                partner_country
            )
        )

        if not user_iso or not partner_iso:

            self.time_zone_info = (
                "국가명을 인식할 수 없어 "
                "자동 계산이 불가능합니다.\n"
                f"사용자 국가: {user_country}\n"
                f"파트너 국가: {partner_country}"
            )

            print(
                "[WARN] 국가명 인식 실패"
            )

            return

        (
            user_tz,
            user_tz_count
        ) = (
            self.tz_resolver
            .get_primary_timezone(
                user_iso
            )
        )

        (
            partner_tz,
            partner_tz_count
        ) = (
            self.tz_resolver
            .get_primary_timezone(
                partner_iso
            )
        )

        if not user_tz or not partner_tz:

            self.time_zone_info = (
                "시간대 정보를 찾을 수 없습니다.\n"
                f"사용자 국가: {user_country}\n"
                f"파트너 국가: {partner_country}"
            )

            print(
                "[WARN] TimeZone 조회 실패"
            )

            return

        user_offset = (
            self.tz_resolver
            .get_utc_offset_hours(
                user_tz
            )
        )

        partner_offset = (
            self.tz_resolver
            .get_utc_offset_hours(
                partner_tz
            )
        )

        # 사용자 시간 - 파트너 시간
        diff = (
            user_offset
            - partner_offset
        )

        # 파트너 근무시간을
        # 사용자 현지 시간으로 변환
        partner_start_in_user_local = (
            self.BUSINESS_START
            + diff
        )

        partner_end_in_user_local = (
            self.BUSINESS_END
            + diff
        )

        overlap_start = max(
            self.BUSINESS_START,
            partner_start_in_user_local
        )

        overlap_end = min(
            self.BUSINESS_END,
            partner_end_in_user_local
        )

        # 겹치는 시간 존재
        if overlap_start < overlap_end:

            overlap_text = (
                f"{user_country} 시간 기준 "
                f"{self.format_hour(overlap_start)}"
                " ~ "
                f"{self.format_hour(overlap_end)}"
            )

            middle_time = (
                overlap_start
                + overlap_end
            ) / 2

            recommended_text = (
                f"{user_country} 시간 기준 "
                f"{self.format_hour(middle_time)}"
            )

        else:

            overlap_text = (
                "정규 근무시간(09~18시) 내 "
                "겹치는 시간대 없음"
            )

            recommended_text = (
                "이른 아침 또는 늦은 저녁 등 "
                "유연근무 시간대 조율 필요"
            )

        # -----------------------------------------------------
        # 다중 TimeZone 국가 경고
        # -----------------------------------------------------
        notes = []

        if user_tz_count > 1:

            notes.append(
                f"{user_country}은(는) "
                f"{user_tz_count}개의 시간대를 보유하므로 "
                f"현재 대표 시간대 {user_tz} 기준으로 계산"
            )

        if partner_tz_count > 1:

            notes.append(
                f"{partner_country}은(는) "
                f"{partner_tz_count}개의 시간대를 보유하므로 "
                f"현재 대표 시간대 {partner_tz} 기준으로 계산"
            )

        multi_tz_note = ""

        if notes:

            multi_tz_note = (
                "\n주의: "
                + " / ".join(notes)
            )

        # -----------------------------------------------------
        # 누가 빠른지 표시
        # -----------------------------------------------------
        if diff > 0:

            diff_direction = (
                "사용자가 빠름"
            )

        elif diff < 0:

            diff_direction = (
                "파트너가 빠름"
            )

        else:

            diff_direction = (
                "동일"
            )

        diff_text = (
            self.format_difference(
                diff
            )
        )

        self.time_zone_info = (

            f"{user_country}({user_tz}) 기준 "
            f"{partner_country}({partner_tz})와의 시차: "
            f"{diff_text}시간 "
            f"({diff_direction})\n"

            f"겹치는 근무시간대: "
            f"{overlap_text}\n"

            f"권장 회의 시간: "
            f"{recommended_text}\n"

            "(현재 시점 기준 계산 / "
            "서머타임 자동 반영)"
            f"{multi_tz_note}"
        )

        print(
            "[INFO] 시차 정보 자동 계산 완료"
        )

        print(
            self.time_zone_info
        )

    # ---------------------------------------------------------
    # OpenAI 분석
    # ---------------------------------------------------------
    def analyze_file(self):

        try:

            if not self.content:

                print(
                    "[ERROR] 분석할 파일 내용이 없습니다."
                )

                return False

            self.calculate_time_zone_info()

            user_message = (

                f"{self.content}\n\n"

                "[시차 정보 - 자동 계산됨]\n"

                f"{self.time_zone_info}"
            )

            print(
                "[INFO] OpenAI API 분석 시작 "
                "(1단계: 정리·요약)"
            )

            response = (
                self.client
                .chat
                .completions
                .create(
                    model=self.model_name,

                    messages=[
                        {
                            "role": "system",
                            "content": self.SYSTEM_PROMPT
                        },
                        {
                            "role": "user",
                            "content": user_message
                        }
                    ],

                    temperature=0.0
                )
            )

            self.result = (
                response
                .choices[0]
                .message
                .content
                or ""
            ).strip()

            if not self.result:

                print(
                    "[ERROR] OpenAI 응답이 비어 있습니다."
                )

                return False

            print(
                "[INFO] OpenAI API 분석 완료"
            )

            return True

        except Exception as e:

            print(
                f"[ERROR] OpenAI API 분석 실패: {e}"
            )

            return False

    # ---------------------------------------------------------
    # JSON 안전 파싱
    # ---------------------------------------------------------
    @staticmethod
    def safe_parse_json(text):

        if not text:
            return None

        text = text.strip()

        # ```json 제거
        if text.startswith("```"):

            text = re.sub(
                r"^```(?:json)?\s*",
                "",
                text,
                flags=re.IGNORECASE
            )

            text = re.sub(
                r"\s*```$",
                "",
                text
            )

        try:

            return json.loads(text)

        except json.JSONDecodeError:

            start = text.find("{")
            end = text.rfind("}")

            if start == -1 or end == -1:
                return None

            try:

                return json.loads(
                    text[start:end + 1]
                )

            except json.JSONDecodeError:

                return None

    # ---------------------------------------------------------
    # 결과 저장
    # ---------------------------------------------------------
    def save_result(self):

        try:

            if not self.result:

                print(
                    "[ERROR] 저장할 결과가 없습니다."
                )

                return False

            parsed_result = (
                self.safe_parse_json(
                    self.result
                )
            )

            if parsed_result is None:

                print(
                    "[WARN] JSON 파싱 실패 - "
                    "raw_result 형태로 저장합니다."
                )

                data = {
                    "raw_result": self.result
                }

            else:

                data = parsed_result

            with open(
                self.result_path,
                "w",
                encoding="utf-8"
            ) as f:

                json.dump(
                    data,
                    f,
                    ensure_ascii=False,
                    indent=2
                )

            print(
                "[INFO] 1단계 결과 저장 완료:"
            )

            print(
                self.result_path
            )

            return True

        except Exception as e:

            print(
                f"[ERROR] 결과 저장 실패: {e}"
            )

            return False

    # ---------------------------------------------------------
    # 저장 결과 출력
    # ---------------------------------------------------------
    def open_result_file(self):

        try:

            if not self.result_path.is_file():

                print(
                    f"[ERROR] 결과 파일 없음: "
                    f"{self.result_path}"
                )

                return

            with open(
                self.result_path,
                "r",
                encoding="utf-8"
            ) as f:

                result_data = f.read()

            print(
                "\n========== "
                "1단계 분석 결과 "
                "=========="
            )

            print(
                result_data
            )

            print(
                "================================="
            )

        except Exception as e:

            print(
                f"[ERROR] 결과 파일 읽기 실패: {e}"
            )


# =============================================================
# MAIN
# =============================================================
if __name__ == "__main__":

    # ---------------------------------------------------------
    # 실행 방법
    #
    # 1. 기본 파일
    # python3 stage1_analyzer.py
    #
    # -> 같은 폴더 meeting_file.txt 사용
    #
    # 2. 다른 파일
    # python3 stage1_analyzer.py /root/files/meeting.pdf
    # ---------------------------------------------------------

    input_file = None

    if len(sys.argv) >= 2:
        input_file = sys.argv[1]

    try:

        analyzer = OpenAIFileAnalyzer(
            input_file_path=input_file
        )

        analyzer.clear_previous_result()

        if analyzer.openfile():

            if analyzer.analyze_file():

                analyzer.save_result()
                analyzer.open_result_file()

    except Exception as e:

        print(
            f"[FATAL ERROR] {e}"
        )
